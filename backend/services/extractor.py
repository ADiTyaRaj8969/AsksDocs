"""
Server-side extractor — kept for the Python backend only.
In normal operation the Node.js server accepts client-side chunks, so this
file is only used if you run the FastAPI backend directly.
"""

import os
import io
from pathlib import Path
from typing import List, Dict

from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from PIL import Image
from google import genai

VISION_MODEL = "gemini-2.5-flash"
_ENV_PATH = Path(__file__).parent.parent / ".env"


def _key() -> str:
    load_dotenv(dotenv_path=_ENV_PATH, override=True)
    return os.environ["GEMINI_API_KEY"]


def _ocr_image(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    image_bytes = buf.getvalue()
    with genai.Client(api_key=_key()) as client:
        response = client.models.generate_content(
            model=VISION_MODEL,
            contents=[
                {"mime_type": "image/png", "data": image_bytes},
                "Extract all text from this image exactly as it appears. "
                "Preserve paragraph structure. Return only the extracted text.",
            ],
        )
    return response.text.strip() if response.text else ""


def extract_text_from_pdf(file_path: str) -> List[Dict]:
    pages = []
    doc = fitz.open(file_path)
    doc_name = os.path.basename(file_path)
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text().strip()
        if text:
            pages.append({"text": text, "page_number": page_num, "document_name": doc_name})
    doc.close()
    return pages


def extract_scanned_pdf(file_path: str) -> List[Dict]:
    pages = []
    doc = fitz.open(file_path)
    doc_name = os.path.basename(file_path)
    for page_num, page in enumerate(doc, start=1):
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = _ocr_image(img)
        if text:
            pages.append({"text": text, "page_number": page_num, "document_name": doc_name})
    doc.close()
    return pages


def extract_text_from_docx(file_path: str) -> List[Dict]:
    doc = Document(file_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Include table cell text (previously missing)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts)
    if not text:
        return []
    return [{"text": text, "page_number": 1, "document_name": os.path.basename(file_path)}]


def extract_text_from_excel(file_path: str) -> List[Dict]:
    # Use sheet index as page_number (1-based) so multi-sheet docs cite correctly
    sheets = pd.read_excel(file_path, sheet_name=None)
    pages = []
    doc_name = os.path.basename(file_path)
    for sheet_index, (sheet_name, df) in enumerate(sheets.items(), start=1):
        text = f"Sheet: {sheet_name}\n\n{df.to_string(index=False)}"
        pages.append({"text": text, "page_number": sheet_index, "document_name": doc_name})
    return pages


def extract_text_from_image(file_path: str) -> List[Dict]:
    img = Image.open(file_path)
    text = _ocr_image(img)
    return [{"text": text, "page_number": 1, "document_name": os.path.basename(file_path)}]


def extract(file_path: str) -> List[Dict]:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        pages = extract_text_from_pdf(file_path)
        if len(" ".join(p["text"] for p in pages).strip()) < 100:
            pages = extract_scanned_pdf(file_path)
        return pages

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    if ext in {".xlsx", ".xls"}:
        return extract_text_from_excel(file_path)

    if ext in {".png", ".jpg", ".jpeg"}:
        return extract_text_from_image(file_path)

    raise ValueError(f"Unsupported file type: {ext}")
