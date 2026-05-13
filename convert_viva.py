"""
Convert VIVA_QA.md to VIVA_QA.docx using python-docx.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = "VIVA_QA.md"
DOCX_PATH = "VIVA_QA.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BG      = RGBColor(0x1E, 0x1E, 0x2E)   # code block bg
CODE_FG      = RGBColor(0xD0, 0xD0, 0xD0)   # code text
H1_COLOR     = RGBColor(0x2E, 0x4D, 0x9E)   # deep blue
H2_COLOR     = RGBColor(0x1F, 0x6F, 0x8B)   # teal-blue
H3_COLOR     = RGBColor(0x2E, 0x7D, 0x32)   # green
QA_Q_COLOR   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black for Q label
QA_A_COLOR   = RGBColor(0x23, 0x5A, 0x4E)   # dark green for A label
BORDER_COLOR = "999999"


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)
        borders.append(el)
    tblPr.append(borders)


def inline_run(para, text, bold=False, italic=False, code=False, color=None):
    """Add a run with optional formatting to a paragraph."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if code:
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
    if color:
        run.font.color.rgb = color
    return run


# Regex for inline markdown tokens
_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*"        # ***bold-italic***
    r"|\*\*(.+?)\*\*"             # **bold**
    r"|\*(.+?)\*"                 # *italic*
    r"|`(.+?)`"                   # `code`
    r"|__(.*?)__"                 # __bold__
    r"|_(.+?)_)",                 # _italic_
    re.DOTALL,
)


def add_inline(para, text):
    """Parse inline markdown and add formatted runs to *para*."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            inline_run(para, text[pos:m.start()])
        raw = m.group(0)
        if raw.startswith("***"):
            inline_run(para, m.group(2), bold=True, italic=True)
        elif raw.startswith("**"):
            inline_run(para, m.group(3), bold=True)
        elif raw.startswith("__"):
            inline_run(para, m.group(6), bold=True)
        elif raw.startswith("*"):
            inline_run(para, m.group(4), italic=True)
        elif raw.startswith("_"):
            inline_run(para, m.group(7), italic=True)
        elif raw.startswith("`"):
            inline_run(para, m.group(5), code=True)
        pos = m.end()
    if pos < len(text):
        inline_run(para, text[pos:])


def add_code_block(doc, lines):
    """Add a shaded code block paragraph."""
    code_text = "\n".join(lines)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # shade the paragraph background via XML
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    run = para.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def parse_table(doc, raw_rows):
    """
    raw_rows is a list of markdown table lines including the separator row.
    """
    data_rows = [r for r in raw_rows if not re.match(r"^\s*\|[-| :]+\|\s*$", r)]
    if not data_rows:
        return

    def split_row(line):
        parts = line.strip().strip("|").split("|")
        return [p.strip() for p in parts]

    parsed  = [split_row(r) for r in data_rows]
    ncols   = max(len(r) for r in parsed)

    table = doc.add_table(rows=len(parsed), cols=ncols)
    add_table_borders(table)

    for ri, row_data in enumerate(parsed):
        row = table.rows[ri]
        for ci in range(ncols):
            cell = row.cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            text = row_data[ci] if ci < len(row_data) else ""
            if ri == 0:
                set_cell_bg(cell, "D6E4F0")
                r = para.add_run(text)
                r.bold = True
                r.font.size = Pt(9.5)
            else:
                add_inline(para, text)
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()  # spacing after table


def build_doc():
    doc = Document()

    # ── Default styles ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for hname, sz, bold, color in [
        ("Heading 1", 18, True,  H1_COLOR),
        ("Heading 2", 14, True,  H2_COLOR),
        ("Heading 3", 12, True,  H3_COLOR),
    ]:
        h = doc.styles[hname]
        h.font.size  = Pt(sz)
        h.font.bold  = bold
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(4)

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Read markdown ──────────────────────────────────────────────────────────
    with open(MD_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    total = len(lines)

    while i < total:
        line = lines[i].rstrip("\n")

        # ── Skip horizontal rules ──────────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        # ── Headings ───────────────────────────────────────────────────────────
        hm = re.match(r"^(#{1,6})\s+(.*)", line)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2).strip()
            # strip anchor links like {#foo}
            text  = re.sub(r"\{#[\w-]+\}", "", text).strip()
            hlevel = min(level, 3)
            doc.add_heading(text, level=hlevel)
            i += 1
            continue

        # ── Code block ────────────────────────────────────────────────────────
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < total and not lines[i].rstrip("\n").startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < total and lines[i].rstrip("\n").startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            parse_table(doc, table_lines)
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Inches(0.4)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Unordered list ────────────────────────────────────────────────────
        if re.match(r"^(\s*)[-*+]\s+", line):
            indent = len(line) - len(line.lstrip())
            text   = re.sub(r"^\s*[-*+]\s+", "", line)
            para   = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Inches(0.3 + indent * 0.02)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline(para, text)
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline(para, text)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Plain paragraph ───────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(4)
        add_inline(para, line)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    build_doc()
