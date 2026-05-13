#!/usr/bin/env python3
"""
generate_srs.py
Generates docs/SRS.docx — Software Requirements Specification for AskDocs.

Run from the project root:
    python generate_srs.py

Requirements (already in backend/requirements.txt):
    pip install python-docx
    pip install matplotlib        # optional — enables flowchart images
"""

import os
import sys

# ── Dependency guard ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed.\n  pip install python-docx")

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, Polygon
    HAS_MPL = True
except ImportError:
    HAS_MPL = False
    print("WARNING: matplotlib not found — flowchart images will be skipped.")
    print("  Install with:  pip install matplotlib\n")

# ── Paths & constants ─────────────────────────────────────────────────────────
ROOT     = os.path.dirname(os.path.abspath(__file__))
OUT_DIR  = os.path.join(ROOT, "docs")
SRS_PATH = os.path.join(OUT_DIR, "SRS.docx")
TODAY    = "April 06, 2026"

os.makedirs(OUT_DIR, exist_ok=True)

# ── Colour palette ────────────────────────────────────────────────────────────
C_PRIMARY = RGBColor(0x43, 0x38, 0xCA)   # Indigo  #4338CA
C_DEEP    = RGBColor(0x1E, 0x1B, 0x4B)   # Deep indigo
C_TEXT    = RGBColor(0x1F, 0x29, 0x37)   # Dark slate
C_MUTED   = RGBColor(0x6B, 0x72, 0x80)   # Gray-500
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# ══════════════════════════════════════════════════════════════════════════════
#  FLOWCHART HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, text, color, text_color="white", fontsize=9, radius=0.06):
    box = FancyBboxPatch(
        (x - w / 2, y - h / 2), w, h,
        boxstyle=f"round,pad=0,rounding_size={radius}",
        linewidth=0, facecolor=color, zorder=3,
    )
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center",
            fontsize=fontsize, color=text_color,
            fontweight="bold", zorder=4, multialignment="center")


def _arrow(ax, x1, y1, x2, y2, color="#94A3B8"):
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=1.6),
        zorder=2,
    )


def _label(ax, x, y, text, color="#475569", size=7.5):
    ax.text(x, y, text, ha="center", va="center",
            fontsize=size, color=color, fontstyle="italic", zorder=5)


# ── Flowchart 1: Document Upload Pipeline ────────────────────────────────────
def make_upload_flow(path):
    fig, ax = plt.subplots(figsize=(7.5, 13))
    ax.set_xlim(0, 7.5)
    ax.set_ylim(0, 13)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    INDIGO  = "#4338CA"
    VIOLET  = "#7C3AED"
    TEAL    = "#0D9488"
    BLUE    = "#0284C7"
    AMBER   = "#D97706"
    EMERALD = "#059669"
    ROSE    = "#E11D48"
    GRAY    = "#64748B"

    # Title
    ax.text(3.75, 12.55, "Document Upload Pipeline", ha="center",
            fontsize=13, fontweight="bold", color="#1E1B4B")

    steps = [
        (3.75, 11.7,  "User Selects File\n(Drag-and-Drop or Browse)",    INDIGO),
        (3.75, 10.5,  "Format Validation\nPDF / DOCX / XLSX / PNG / JPG", VIOLET),
        (3.75,  9.3,  "Text Extraction",                                   VIOLET),
        (3.75,  8.1,  "Is Scanned / Image?\n(< 100 chars extracted)",     AMBER),
        (3.75,  6.9,  "Gemini Vision OCR\n(Page-by-page rendering)",      AMBER),
        (3.75,  5.7,  "Sliding-Window Chunking\n500 words · 100 overlap", INDIGO),
        (3.75,  4.5,  "Batch Embedding\nGemini text-embedding-004",        BLUE),
        (3.75,  3.3,  "Store Chunks + Vectors\nChromaDB (Cosine HNSW)",   TEAL),
        (3.75,  2.1,  "Return UploadResponse\n{document_name, chunks_created}", EMERALD),
    ]

    for x, y, label, color in steps:
        _box(ax, x, y, 3.6, 0.78, label, color, fontsize=8.5)

    # Side box: extraction libraries
    _box(ax, 0.95, 9.3, 1.5, 0.78,
         "PyMuPDF\npython-docx\nPandas", GRAY, fontsize=7.5)
    ax.annotate("", xy=(2.25, 9.3 - 0.01), xytext=(1.7, 9.3),
                arrowprops=dict(arrowstyle="-|>", color=GRAY, lw=1.2))

    # Side box: rejection
    _box(ax, 0.85, 10.5, 1.4, 0.6, "HTTP 400\nRejected", ROSE, fontsize=8)
    ax.annotate("", xy=(1.55, 10.5), xytext=(1.95, 10.5),
                arrowprops=dict(arrowstyle="-|>", color=ROSE, lw=1.2))
    _label(ax, 1.75, 10.65, "Invalid")

    # Arrows between main steps
    skip = {3}   # step 3 (scanned?) branches to OCR or directly to chunking
    for i in range(len(steps) - 1):
        if i in skip:
            continue
        _, y1, _, _ = steps[i]
        _, y2, _, _ = steps[i + 1]
        _arrow(ax, 3.75, y1 - 0.39, 3.75, y2 + 0.39)

    # Arrow from scanned check -> OCR (Yes)
    _arrow(ax, 3.75, 8.1 - 0.39, 3.75, 6.9 + 0.39)
    _label(ax, 4.05, 7.5, "Yes (OCR)")

    # Arrow bypass: scanned=No -> chunking  (right side bypass)
    ax.annotate("", xy=(5.8, 5.7), xytext=(5.8, 8.1),
                arrowprops=dict(arrowstyle="-|>", color="#0EA5E9", lw=1.4,
                                connectionstyle="arc3,rad=0.0"))
    ax.annotate("", xy=(3.75 + 1.8, 5.7), xytext=(5.8, 5.7),
                arrowprops=dict(arrowstyle="-|>", color="#0EA5E9", lw=1.4))
    _label(ax, 6.3, 7.0, "No\n(Skip OCR)", color="#0EA5E9")

    # Remaining arrows after OCR
    for i in [4, 5, 6, 7]:
        _, y1, _, _ = steps[i]
        _, y2, _, _ = steps[i + 1]
        _arrow(ax, 3.75, y1 - 0.39, 3.75, y2 + 0.39)

    plt.tight_layout(pad=0.4)
    plt.savefig(path, dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close()
    print(f"  -> {path}")


# ── Flowchart 2: Query Processing Pipeline ───────────────────────────────────
def make_query_flow(path):
    fig, ax = plt.subplots(figsize=(7.5, 12))
    ax.set_xlim(0, 7.5)
    ax.set_ylim(0, 12)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    INDIGO  = "#4338CA"
    VIOLET  = "#7C3AED"
    BLUE    = "#0284C7"
    TEAL    = "#0D9488"
    EMERALD = "#059669"
    AMBER   = "#D97706"
    ROSE    = "#E11D48"

    ax.text(3.75, 11.55, "Query Processing Pipeline", ha="center",
            fontsize=13, fontweight="bold", color="#1E1B4B")

    steps = [
        (3.75, 10.8, "User Submits Question",                              INDIGO),
        (3.75,  9.7, "Validate: non-empty input",                          VIOLET),
        (3.75,  8.6, "Embed Query\nGemini text-embedding-004",             BLUE),
        (3.75,  7.5, "Cosine Similarity Search\nChromaDB — top_k = 5",   TEAL),
        (3.75,  6.4, "Relevant Chunks Found?",                             AMBER),
        (3.75,  5.3, "Build Grounded Prompt\nContext + Question injection", INDIGO),
        (3.75,  4.2, "Gemini 2.5 Flash\nAnswer Generation",               VIOLET),
        (3.75,  3.1, "Deduplicate Citations\n(document_name, page_number)", BLUE),
        (3.75,  2.0, "Return Answer + Citations\nto React Frontend",       EMERALD),
    ]

    for x, y, label, color in steps:
        _box(ax, x, y, 3.6, 0.78, label, color, fontsize=8.5)

    # No-chunks branch
    _box(ax, 6.4, 6.4, 1.8, 0.68,
         '"No documents\nuploaded"', ROSE, fontsize=8)
    ax.annotate("", xy=(5.55, 6.4), xytext=(5.5, 6.4),
                arrowprops=dict(arrowstyle="-|>", color=ROSE, lw=1.2))
    _label(ax, 5.25, 6.58, "No")

    # All vertical arrows
    for i in range(len(steps) - 1):
        _, y1, _, _ = steps[i]
        _, y2, _, _ = steps[i + 1]
        _arrow(ax, 3.75, y1 - 0.39, 3.75, y2 + 0.39)

    # "Yes" label on arrow from found? -> build prompt
    _label(ax, 4.05, 5.85, "Yes")

    plt.tight_layout(pad=0.4)
    plt.savefig(path, dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close()
    print(f"  -> {path}")


# ── Flowchart 3: System Architecture ─────────────────────────────────────────
def make_architecture(path):
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    fig.patch.set_facecolor("#EEF2FF")
    ax.set_facecolor("#EEF2FF")

    INDIGO  = "#4338CA"
    VIOLET  = "#6D28D9"
    BLUE    = "#0284C7"
    TEAL    = "#0D9488"
    EMERALD = "#059669"
    AMBER   = "#B45309"
    ROSE    = "#BE123C"
    GRAY    = "#475569"

    ax.text(6.5, 7.15, "AskDocs — System Architecture Overview",
            ha="center", fontsize=14, fontweight="bold", color="#1E1B4B")

    # ── Frontend block ────────────────────────────────────────────────────────
    _box(ax, 2.2, 6.1, 3.8, 0.85, "React Frontend  (Vite + Tailwind CSS)", INDIGO, fontsize=10)
    comps = ["FileUpload", "DocumentList", "ChatInterface", "CitationPanel"]
    for i, c in enumerate(comps):
        _box(ax, 0.55 + i * 1.0, 5.1, 0.88, 0.6, c, VIOLET, fontsize=7.5)

    # ── Backend block ─────────────────────────────────────────────────────────
    _box(ax, 9.3, 6.1, 5.2, 0.85, "FastAPI Backend  (Python 3.10 / Uvicorn)", BLUE, fontsize=10)
    endpoints = ["POST\n/upload", "GET\n/documents", "DELETE\n/documents", "POST\n/query"]
    ep_x = [7.3, 8.5, 9.7, 10.9]
    for i, (ep, x) in enumerate(zip(endpoints, ep_x)):
        _box(ax, x, 5.1, 1.0, 0.6, ep, "#0369A1", fontsize=7.5)

    # ── Services layer ────────────────────────────────────────────────────────
    ax.text(6.5, 4.35, "Backend Services", ha="center", fontsize=9,
            color=GRAY, fontstyle="italic")
    services = [
        (1.3,  3.75, "Extractor\nPyMuPDF\n+ OCR",    TEAL),
        (2.9,  3.75, "Chunker\n500w /\n100 overlap",  TEAL),
        (4.5,  3.75, "Embedder\nGemini\ntext-004",    TEAL),
        (6.1,  3.75, "Vector\nStore\nChromaDB",       TEAL),
        (7.7,  3.75, "LLM\nGemini\n2.5 Flash",        VIOLET),
        (9.3,  3.75, "Citation\nDedup\n+ Format",     EMERALD),
    ]
    for x, y, label, color in services:
        _box(ax, x, y, 1.4, 0.9, label, color, fontsize=8)

    # ── Data / external layer ──────────────────────────────────────────────────
    _box(ax, 3.4,  2.1, 3.2, 0.85,
         "Google Gemini API\nEmbeddings · Vision OCR · LLM", ROSE, fontsize=9)
    _box(ax, 8.5,  2.1, 3.2, 0.85,
         "ChromaDB\nLocal Persistent Vector DB (HNSW)", AMBER, fontsize=9)

    # ── Arrows ────────────────────────────────────────────────────────────────
    # Frontend ↔ Backend
    _arrow(ax, 4.1, 6.1, 6.7, 6.1, color="#94A3B8")
    ax.text(5.4, 6.22, "HTTP / REST  (localhost:8000)",
            ha="center", fontsize=7.5, color=GRAY)

    # Backend -> Services
    for sx, sy, _, _ in services:
        _arrow(ax, 9.3, 5.67, sx, sy + 0.45, color="#94A3B8")

    # Services -> External
    for sx, sy, _, _ in services[:3]:      # extractor, chunker, embedder -> Gemini
        _arrow(ax, sx, sy - 0.45, 3.4, 2.53, color="#94A3B8")
    _arrow(ax, 6.1, 3.75 - 0.45, 8.5, 2.53, color="#94A3B8")   # vector store -> ChromaDB
    _arrow(ax, 7.7, 3.75 - 0.45, 4.5, 2.53, color="#94A3B8")   # LLM -> Gemini

    # Legend
    legend = [("#4338CA","Frontend"), ("#0284C7","Backend API"),
              ("#0D9488","Services"), ("#BE123C","External APIs")]
    for i, (color, label) in enumerate(legend):
        _box(ax, 1.0 + i * 2.9, 0.5, 2.2, 0.55, label, color, fontsize=8.5)

    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close()
    print(f"  -> {path}")


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False, italic=False, size=10.5,
             color=None, indent_cm=0):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold            = True
                run.font.size       = Pt(9.5)
                run.font.color.rgb  = C_WHITE
        _set_cell_shade(cell, "4338CA")

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                _set_cell_shade(cell, "EEF2FF")

    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])

    return table


def add_hr(doc, color_hex="4338CA"):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    for run in p.runs:
        run.font.size      = Pt(9)
        run.italic         = True
        run.font.color.rgb = C_MUTED


def embed_image(doc, path, width_inches=6.0, caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_srs():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name       = "Calibri"
    normal.font.size       = Pt(10.5)
    normal.font.color.rgb  = C_TEXT

    # Heading styles
    for name, size, hex_col in [
        ("Heading 1", 16, "4338CA"),
        ("Heading 2", 13, "1E1B4B"),
        ("Heading 3", 11, "374151"),
    ]:
        h = doc.styles[name]
        h.font.name      = "Calibri"
        h.font.bold      = True
        h.font.color.rgb = RGBColor.from_string(hex_col)
        h.font.size      = Pt(size)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AskDocs")
    r.bold = True; r.font.size = Pt(42); r.font.color.rgb = C_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Intelligent Document Query System")
    r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification (SRS)")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = C_TEXT

    doc.add_paragraph()
    add_hr(doc)
    doc.add_paragraph()

    meta = [
        ("Document Version", "1.0"),
        ("Date",             TODAY),
        ("Course",           "Advanced Web Technology (AWT)"),
        ("Semester",         "6th Semester — B.Tech"),
        ("Project Type",     "Full-Stack RAG Web Application"),
        ("Technology",       "FastAPI · React · ChromaDB · Google Gemini"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}:   ")
        r1.bold = True; r1.font.size = Pt(11.5); r1.font.color.rgb = C_MUTED
        r2 = p.add_run(value)
        r2.font.size = Pt(11.5); r2.font.color.rgb = C_TEXT
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", 1)
    add_hr(doc)

    toc = [
        ("1.",   "Introduction",                                  False),
        ("1.1",  "Purpose",                                       True),
        ("1.2",  "Scope",                                         True),
        ("1.3",  "Definitions, Acronyms, and Abbreviations",      True),
        ("1.4",  "Document Overview",                             True),
        ("2.",   "Overall Description",                           False),
        ("2.1",  "Product Perspective",                           True),
        ("2.2",  "Product Functions",                             True),
        ("2.3",  "User Classes and Characteristics",              True),
        ("2.4",  "Operating Environment",                         True),
        ("2.5",  "Constraints and Assumptions",                   True),
        ("3.",   "System Architecture",                           False),
        ("3.1",  "High-Level Architecture Diagram",               True),
        ("3.2",  "Document Upload Pipeline",                      True),
        ("3.3",  "Query Processing Pipeline",                     True),
        ("4.",   "Functional Requirements",                       False),
        ("4.1",  "Document Management",                           True),
        ("4.2",  "Text Extraction",                               True),
        ("4.3",  "Chunking and Embedding",                        True),
        ("4.4",  "Vector Search",                                 True),
        ("4.5",  "Answer Generation",                             True),
        ("5.",   "Non-Functional Requirements",                   False),
        ("6.",   "External Interface Requirements (REST API)",    False),
        ("6.1",  "API Endpoints",                                 True),
        ("6.2",  "Data Schemas",                                  True),
        ("6.3",  "Frontend–Backend Communication",                True),
        ("7.",   "Technology Stack",                              False),
        ("8.",   "System Constraints",                            False),
        ("8.1",  "Technical Constraints",                         True),
        ("8.2",  "Business Constraints",                          True),
        ("8.3",  "Regulatory Constraints",                        True),
    ]

    for num, text, is_sub in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.8) if is_sub else Cm(0)
        run = p.add_run(f"{num}   {text}")
        run.font.size      = Pt(10.5)
        run.font.color.rgb = C_TEXT
        if not is_sub:
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Introduction", 1)
    add_hr(doc)

    add_heading(doc, "1.1 Purpose", 2)
    add_para(doc,
        "This Software Requirements Specification (SRS) describes the complete functional and "
        "non-functional requirements for AskDocs — a privacy-first Intelligent Document Query "
        "System built on Retrieval-Augmented Generation (RAG). The document is the primary "
        "reference for developers, evaluators, and academic assessors reviewing this deliverable "
        "for the 6th Semester Advanced Web Technology course."
    )

    add_heading(doc, "1.2 Scope", 2)
    add_para(doc,
        "AskDocs enables users to upload documents in multiple formats (PDF, DOCX, XLSX, PNG, "
        "JPG) and query their content using natural language. The system extracts, chunks, embeds, "
        "and stores document content locally using ChromaDB, then answers questions by retrieving "
        "the most relevant text passages and generating grounded, cited responses via the Google "
        "Gemini API."
    )
    add_para(doc, "Key scope boundaries:", bold=True)
    for item in [
        "All document data persists locally — only minimal retrieved context is transmitted to Gemini.",
        "Designed as a single-user local deployment; no cloud hosting or authentication is included.",
        "Multi-document queries are supported — answers may draw from several uploaded files.",
        "Out of scope: multi-user authentication, real-time collaboration, and cloud deployment.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Definitions, Acronyms, and Abbreviations", 2)
    add_table(doc,
        ["Term", "Definition"],
        [
            ["RAG",              "Retrieval-Augmented Generation — LLM answers grounded in retrieved document context"],
            ["IDP",              "Intelligent Document Processing — automated extraction and understanding of documents"],
            ["ChromaDB",         "Open-source, embedded vector database for storing and querying text embeddings"],
            ["Embedding",        "A fixed-length numerical vector representing the semantic meaning of a text passage"],
            ["Chunk",            "A fixed-size segment of extracted document text (500 words, 100-word overlap)"],
            ["top_k",            "Number of most-similar chunks retrieved per query (default: 5)"],
            ["HNSW",             "Hierarchical Navigable Small World — fast approximate nearest-neighbour algorithm"],
            ["OCR",              "Optical Character Recognition — converts image-based text to machine-readable text"],
            ["LLM",              "Large Language Model — generates natural-language answers (Gemini 2.5 Flash)"],
            ["REST API",         "HTTP interface between the React frontend and the FastAPI backend"],
            ["FastAPI",          "Modern async Python framework for building REST APIs with OpenAPI docs"],
            ["Vite",             "Frontend build tool and dev server for the React SPA"],
            ["Cosine Similarity","Metric measuring the angle between two vectors to determine semantic closeness"],
        ],
        col_widths=[1.6, 5.0],
    )

    add_heading(doc, "1.4 Document Overview", 2)
    add_para(doc,
        "Section 2 provides a high-level product description. Section 3 presents the system "
        "architecture with three flowcharts. Sections 4 and 5 detail functional and non-functional "
        "requirements. Section 6 specifies the REST API. Sections 7 and 8 cover the technology "
        "stack and system constraints."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. OVERALL DESCRIPTION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Overall Description", 1)
    add_hr(doc)

    add_heading(doc, "2.1 Product Perspective", 2)
    add_para(doc,
        "AskDocs is a standalone full-stack web application following a three-tier architecture: "
        "a React SPA (presentation), a FastAPI REST server (application logic), and a local "
        "ChromaDB instance (data persistence). The React frontend communicates exclusively with "
        "the FastAPI backend through HTTP REST calls. The backend orchestrates a complete RAG "
        "pipeline: extract -> chunk -> embed -> store -> retrieve -> generate. The only external "
        "dependency is the Google Gemini API for embeddings (text-embedding-004) and answer "
        "generation (Gemini 2.5 Flash)."
    )

    add_heading(doc, "2.2 Product Functions", 2)
    for f in [
        "Document Ingestion — accept multi-format file uploads; save to local disk and process",
        "Text Extraction — extract text from PDF, DOCX, Excel, and images (with OCR fallback)",
        "Intelligent Chunking — split extracted text into overlapping 500-word semantic segments",
        "Vector Embedding — convert chunks to 768-dimensional vectors using Gemini text-embedding-004",
        "Semantic Storage — persist vectors and metadata in a local ChromaDB cosine-similarity index",
        "Natural-Language Querying — embed user questions and retrieve the top-5 similar chunks",
        "Grounded Answer Generation — produce cited, Markdown-formatted answers using Gemini 2.5 Flash",
        "Citation Tracking — surface source document and page number for every generated answer",
        "Document Management — list all uploaded documents; delete documents with full vector cleanup",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "2.3 User Classes and Characteristics", 2)
    add_table(doc,
        ["User Class", "Description", "Technical Level"],
        [
            ["End User",   "Uploads personal or professional documents; queries them conversationally", "Non-technical"],
            ["Student",    "Uses AskDocs to search study materials, notes, and research papers",        "Basic"],
            ["Researcher", "Queries large document corpora to extract specific factual information",     "Intermediate"],
            ["Developer",  "Runs the application locally; may extend or modify the codebase",           "Advanced"],
        ],
        col_widths=[1.5, 3.8, 1.3],
    )

    add_heading(doc, "2.4 Operating Environment", 2)
    add_table(doc,
        ["Component", "Requirement"],
        [
            ["Operating System",   "Windows 10/11, macOS 12+, or Ubuntu 20.04+"],
            ["Python",             "3.10 or higher"],
            ["Node.js",            "18.x or higher (for frontend build and dev server)"],
            ["Web Browser",        "Chrome 110+, Firefox 110+, or Edge 110+"],
            ["Network Access",     "Active internet connection required for Gemini API calls"],
            ["Disk Space",         "Minimum 500 MB free for ChromaDB, uploads, and Python venv"],
            ["Gemini API Key",     "Valid Google Gemini API key stored in backend/.env"],
        ],
        col_widths=[2.2, 4.4],
    )

    add_heading(doc, "2.5 Constraints and Assumptions", 2)
    add_para(doc, "Constraints:", bold=True)
    for c in [
        "Gemini API rate limits apply; sustained batch embedding may be throttled.",
        "Very large documents (>50 MB) may cause slow processing or memory pressure.",
        "CORS policy restricts API access to localhost:5173 and localhost:3000.",
        "ChromaDB uses cosine similarity only; Euclidean or dot-product distance is not configured.",
        "No streaming — the frontend waits for the full LLM response before rendering.",
    ]:
        add_bullet(doc, c)

    add_para(doc, "Assumptions:", bold=True)
    for a in [
        "The user supplies a valid GEMINI_API_KEY in backend/.env before starting the server.",
        "The system runs on a trusted local machine — no authentication layer is required.",
        "Uploaded documents are in one of the six supported formats.",
        "Questions are in English; multilingual support is not guaranteed.",
    ]:
        add_bullet(doc, a)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. System Architecture", 1)
    add_hr(doc)

    add_heading(doc, "3.1 High-Level Architecture Diagram", 2)
    add_para(doc,
        "AskDocs follows a three-tier architecture. The React SPA (presentation tier) communicates "
        "with the FastAPI backend (logic tier) over HTTP REST. The backend delegates to six isolated "
        "service modules that interact with the ChromaDB vector store (data tier) and the Google "
        "Gemini API (external AI tier)."
    )

    if HAS_MPL:
        arch_path = os.path.join(OUT_DIR, "architecture.png")
        make_architecture(arch_path)
        embed_image(doc, arch_path, width_inches=6.2,
                    caption="Figure 1 — AskDocs System Architecture Overview")

    add_heading(doc, "3.2 Document Upload Pipeline", 2)
    add_para(doc,
        "When the user uploads a file the system validates the extension, extracts raw text "
        "(falling back to Gemini Vision OCR for scanned PDFs and images), applies overlapping "
        "sliding-window chunking, calls the Gemini embedding API in a single batch, and stores "
        "all vectors and metadata in ChromaDB. Re-uploading the same filename first deletes "
        "the stale chunks, ensuring idempotent behaviour."
    )

    if HAS_MPL:
        upload_path = os.path.join(OUT_DIR, "upload_flow.png")
        make_upload_flow(upload_path)
        embed_image(doc, upload_path, width_inches=4.8,
                    caption="Figure 2 — Document Upload Pipeline Flowchart")

    add_heading(doc, "3.3 Query Processing Pipeline", 2)
    add_para(doc,
        "A user question is embedded using the same Gemini text-embedding-004 model, then used to "
        "perform a cosine similarity search across all stored chunks. The top-5 most relevant "
        "passages are injected into a grounded prompt alongside the original question. Gemini "
        "2.5 Flash generates a cited, Markdown-formatted answer. The backend deduplicates citations "
        "by (document_name, page_number) before returning the response to the frontend."
    )

    if HAS_MPL:
        query_path = os.path.join(OUT_DIR, "query_flow.png")
        make_query_flow(query_path)
        embed_image(doc, query_path, width_inches=4.8,
                    caption="Figure 3 — Query Processing Pipeline Flowchart")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. FUNCTIONAL REQUIREMENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Functional Requirements", 1)
    add_hr(doc)

    req_sections = [
        ("4.1 Document Management", [
            ["FR-01", "Upload",         "System shall accept single-file uploads via drag-and-drop or file browser.",        "High"],
            ["FR-02", "Format Support", "System shall support PDF, DOCX, XLSX, XLS, PNG, JPG, and JPEG file types.",         "High"],
            ["FR-03", "Validation",     "System shall reject unsupported file types with HTTP 400 and a descriptive message.","High"],
            ["FR-04", "Re-upload",      "Re-uploading the same filename shall delete its previous vector chunks first.",      "Medium"],
            ["FR-05", "List Docs",      "System shall provide GET /api/documents returning all document names in ChromaDB.",  "High"],
            ["FR-06", "Delete Doc",     "System shall delete a document and all its vector chunks via DELETE /api/documents/{name}.", "High"],
        ]),
        ("4.2 Text Extraction", [
            ["FR-07", "PDF Native",   "System shall extract selectable text from PDFs using PyMuPDF.",                        "High"],
            ["FR-08", "OCR Fallback", "If <100 chars extracted from a PDF, system shall fall back to Gemini Vision OCR.",    "High"],
            ["FR-09", "DOCX",         "System shall extract paragraph text from Word documents using python-docx.",           "High"],
            ["FR-10", "Excel",        "System shall serialize all spreadsheet sheets, preserving sheet names and structure.", "High"],
            ["FR-11", "Image OCR",    "System shall extract text from PNG/JPG images via Gemini Vision OCR.",                 "High"],
            ["FR-12", "Empty Guard",  "If no text can be extracted, system shall return HTTP 422 with a clear message.",     "Medium"],
        ]),
        ("4.3 Chunking and Embedding", [
            ["FR-13", "Chunk Size",    "System shall produce word-based chunks of maximum 500 words per chunk.",              "High"],
            ["FR-14", "Overlap",       "Adjacent chunks shall overlap by 100 words to preserve cross-boundary context.",     "High"],
            ["FR-15", "Chunk UUID",    "Each chunk shall carry a unique UUID to support idempotent insertion and deletion.",  "High"],
            ["FR-16", "Batch Embed",   "Embeddings shall be generated in a single batch API call to minimise latency.",      "Medium"],
            ["FR-17", "Embed Model",   "Embedding model shall be Gemini text-embedding-004 producing 768-dim vectors.",      "High"],
        ]),
        ("4.4 Vector Search", [
            ["FR-18", "Cosine Search", "Vector retrieval shall use cosine similarity via ChromaDB's HNSW index.",             "High"],
            ["FR-19", "top_k Param",   "Default retrieval shall return 5 chunks; callers may override top_k per request.",   "High"],
            ["FR-20", "No-Doc Guard",  "If no chunks exist, system shall return a friendly 'upload first' message.",          "High"],
        ]),
        ("4.5 Answer Generation", [
            ["FR-21", "Grounded LLM",  "LLM shall use ONLY retrieved context — no outside knowledge or hallucination.",      "High"],
            ["FR-22", "Inline Cite",   "Every answer shall include inline citations: *(Source: filename, Page N)*.",          "High"],
            ["FR-23", "Dedup Cite",    "The citations list shall deduplicate entries by (document_name, page_number).",      "Medium"],
            ["FR-24", "Markdown",      "Answers shall be formatted in Markdown (headers, bold, bullet lists).",              "Medium"],
            ["FR-25", "Not-Found Msg", "If the answer is absent from context, system shall say so rather than fabricate.",   "High"],
        ]),
    ]

    for section_title, rows in req_sections:
        add_heading(doc, section_title, 2)
        add_table(doc,
            ["ID", "Name", "Description", "Priority"],
            rows,
            col_widths=[0.7, 1.1, 4.2, 0.7],
        )
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. NON-FUNCTIONAL REQUIREMENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Non-Functional Requirements", 1)
    add_hr(doc)

    nfr = [
        ["NFR-01", "Performance",     "Document upload and processing shall complete within 30 seconds for files ≤ 10 MB.",                  "High"],
        ["NFR-02", "Performance",     "Query response (embed -> search -> generate) shall complete within 10 seconds.",                       "High"],
        ["NFR-03", "Reliability",     "Backend shall return structured JSON error responses (HTTP 4xx/5xx) for all failure modes.",          "High"],
        ["NFR-04", "Reliability",     "ChromaDB shall persist data across server restarts without any data loss.",                           "High"],
        ["NFR-05", "Usability",       "Upload area shall display real-time progress percentage during processing.",                          "Medium"],
        ["NFR-06", "Usability",       "All user-facing error messages shall be human-readable — no raw Python tracebacks.",                  "High"],
        ["NFR-07", "Privacy",         "Document text shall not be stored on any external server; only API calls leave localhost.",           "High"],
        ["NFR-08", "Privacy",         "The Gemini API key shall be stored only in backend/.env and never exposed to the frontend.",         "High"],
        ["NFR-09", "Scalability",     "ChromaDB shall handle at least 50 documents (≈25,000 chunks) without performance degradation.",      "Medium"],
        ["NFR-10", "Maintainability", "Each backend service (extractor, chunker, embedder, vector_store, llm) shall be fully isolated.",     "High"],
        ["NFR-11", "Portability",     "Application shall run unchanged on Windows 10/11, macOS 12+, and Ubuntu 20.04+.",                    "Medium"],
        ["NFR-12", "Security",        "CORS policy shall restrict API access to localhost origins (ports 5173 and 3000) only.",              "High"],
        ["NFR-13", "Observability",   "Backend shall log all errors with stack traces for developer debugging.",                             "Low"],
    ]
    add_table(doc,
        ["ID", "Category", "Description", "Priority"],
        nfr,
        col_widths=[0.8, 1.3, 4.0, 0.8],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. EXTERNAL INTERFACE REQUIREMENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. External Interface Requirements (REST API)", 1)
    add_hr(doc)

    add_heading(doc, "6.1 API Endpoints", 2)
    add_table(doc,
        ["Method", "Endpoint", "Description", "Request Body", "Success Response"],
        [
            ["POST",   "/api/upload",            "Upload and process a document",   "multipart/form-data: file",      "200 UploadResponse"],
            ["GET",    "/api/documents",          "List all stored documents",       "—",                              "200 DocumentListResponse"],
            ["DELETE", "/api/documents/{name}",   "Delete document + all its chunks","—",                              "200 {message: str}"],
            ["POST",   "/api/query",              "Ask a natural-language question", "QueryRequest {question, top_k}", "200 QueryResponse"],
            ["GET",    "/health",                 "Backend health check",            "—",                              "200 {status: 'ok'}"],
            ["GET",    "/",                       "Root readiness check",            "—",                              "200 {status, message}"],
        ],
        col_widths=[0.7, 1.9, 1.6, 1.7, 1.8],
    )

    add_heading(doc, "6.2 Data Schemas", 2)
    add_table(doc,
        ["Schema", "Fields"],
        [
            ["UploadResponse",       "message: str · document_name: str · chunks_created: int"],
            ["DocumentListResponse", "documents: List[str]"],
            ["QueryRequest",         "question: str · top_k: int = 5"],
            ["Citation",             "document_name: str · page_number: int · chunk_text: str (≤ 200 chars)"],
            ["QueryResponse",        "answer: str (Markdown) · citations: List[Citation]"],
        ],
        col_widths=[2.0, 5.0],
    )

    add_heading(doc, "6.3 Frontend–Backend Communication", 2)
    add_para(doc,
        "The React frontend communicates with the backend using the Axios HTTP client. All API "
        "calls target http://localhost:8000/api. The Vite dev server runs on http://localhost:5173. "
        "CORS is pre-configured in main.py to permit both localhost:5173 and localhost:3000 origins. "
        "All request and response bodies use JSON (Content-Type: application/json) except file "
        "uploads, which use multipart/form-data."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Technology Stack", 1)
    add_hr(doc)

    add_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Frontend",    "React",                   "18.x",    "UI component library (SPA)"],
            ["Frontend",    "Vite",                    "5.x",     "Build tool and HMR dev server"],
            ["Frontend",    "Tailwind CSS",             "3.x",     "Utility-first CSS framework"],
            ["Frontend",    "React Dropzone",           "14.x",    "Drag-and-drop file upload"],
            ["Frontend",    "Axios",                    "1.x",     "HTTP client for REST API calls"],
            ["Backend",     "Python",                   "3.10+",   "Primary backend language"],
            ["Backend",     "FastAPI",                  "0.111+",  "Async REST API framework"],
            ["Backend",     "Uvicorn",                  "0.30+",   "ASGI server for FastAPI"],
            ["Backend",     "Pydantic v2",              "2.7+",    "Request / response schema validation"],
            ["Backend",     "python-dotenv",            "1.0+",    "Environment variable management"],
            ["AI / ML",     "Google Gemini 2.5 Flash",  "—",       "LLM for grounded answer generation"],
            ["AI / ML",     "Gemini text-embedding-004","—",       "768-dimensional semantic embeddings"],
            ["AI / ML",     "Gemini Vision (2.5 Flash)","—",       "OCR for scanned PDFs and images"],
            ["Data",        "ChromaDB",                 "0.5+",    "Local vector DB with cosine HNSW index"],
            ["Extraction",  "PyMuPDF (fitz)",           "1.24+",   "Fast PDF text extraction"],
            ["Extraction",  "python-docx",              "1.1+",    "Word document (.docx) parsing"],
            ["Extraction",  "Pandas + openpyxl",        "2.2+",    "Excel spreadsheet parsing"],
            ["Extraction",  "Pillow",                   "10.0+",   "Image loading and preprocessing"],
        ],
        col_widths=[1.2, 1.9, 0.9, 3.0],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. SYSTEM CONSTRAINTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. System Constraints", 1)
    add_hr(doc)

    add_heading(doc, "8.1 Technical Constraints", 2)
    for c in [
        "Gemini API is subject to Google's usage quotas; sustained high-volume usage may be throttled.",
        "ChromaDB is a single-process embedded database — concurrent write access is not supported.",
        "The system has no user-session isolation; all uploaded documents are accessible to any browser client on the same localhost.",
        "Scanned PDF OCR consumes one Gemini Vision API call per page, which may be cost-intensive for long documents.",
        "Responses are not streamed — the React frontend waits for the full LLM completion before rendering.",
        "The Gemini text-embedding-004 model produces 768-dimensional vectors; this dimension is fixed and cannot be changed without rebuilding the ChromaDB collection.",
    ]:
        add_bullet(doc, c)

    add_heading(doc, "8.2 Business Constraints", 2)
    for c in [
        "AskDocs is an academic deliverable for the 6th Semester Advanced Web Technology course.",
        "All Google Gemini API usage is subject to Google's Terms of Service and acceptable use policy.",
        "The system is not intended for production deployment without additional security hardening (authentication, rate limiting, HTTPS).",
    ]:
        add_bullet(doc, c)

    add_heading(doc, "8.3 Regulatory Constraints", 2)
    for c in [
        "Users are responsible for ensuring they have the legal right to process the documents they upload.",
        "Minimal data is transmitted externally: only retrieved text chunks are sent to Gemini for embedding/generation; no personally identifiable information is stored beyond what the user uploads.",
        "The system stores uploaded files on local disk; users should ensure compliance with any applicable data protection regulations for their use case.",
    ]:
        add_bullet(doc, c)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(SRS_PATH)
    return SRS_PATH


# ══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating AskDocs SRS...")

    if HAS_MPL:
        print("Generating flowchart images:")
    else:
        print("Skipping flowcharts (matplotlib not installed).")

    path = build_srs()
    print(f"\nDone!  SRS saved to:\n  {path}")
    print("\nTip: Open in Microsoft Word and press Ctrl+A -> F9 to update the TOC page numbers.")
